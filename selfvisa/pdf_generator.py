"""
출입국 통합신청서(신고서) PDF 생성기
checkvisa.co.kr /selfvisa 용

사용법:
    from pdf_generator import fill_form
    pdf_bytes = fill_form(data_dict)

TEMPLATE 경로를 실제 서버 경로로 수정하세요.
"""
import os, io, tempfile, subprocess
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfWriter, PdfReader

# ★ 서버 배포 시 이 경로를 실제 경로로 수정하세요 ★
TEMPLATE = os.path.join(os.path.dirname(__file__), 'assets', '통합신청서(신고서).doc')

ROW_HEIGHTS = {
    9: 227, 17: 360, 24: 432, 25: 432,
    26: 194, 31: 625, 37: 133, 45: 102, 46: 111,
}
EXACT_ROWS = {17, 24, 25}


def set_row_height(row, twips_val, exact=False):
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr'); row._tr.insert(0, trPr)
    trH = trPr.find(qn('w:trHeight'))
    if trH is None:
        trH = OxmlElement('w:trHeight'); trPr.append(trH)
    trH.set(qn('w:val'), str(twips_val))
    trH.set(qn('w:hRule'), 'exact' if exact else 'atLeast')


def shrink_cell_english_font(cell, new_sz=14):
    for para in cell.paragraphs:
        for run in para.runs:
            rPr = run._element.get_or_add_rPr()
            for local in ['w:sz', 'w:szCs']:
                e = rPr.find(qn(local))
                if e is None:
                    e = OxmlElement(local); rPr.append(e)
                e.set(qn('w:val'), str(new_sz))


def write_cell(cell, text):
    if not text: return
    para = cell.paragraphs[0]
    for run in list(para.runs): run._element.getparent().remove(run._element)
    run = para.add_run(text)
    run.font.size = Pt(9); run.font.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    # 나눔고딕 또는 Malgun Gothic (Docker 환경에서 나눔고딕 사용)
    font_name = 'NanumGothic'
    for attr in ['w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs']:
        rFonts.set(qn(attr), font_name)
    for local, val in [('w:sz', '18'), ('w:szCs', '18')]:
        e = rPr.find(qn(local))
        if e is None: e = OxmlElement(local); rPr.append(e)
        e.set(qn('w:val'), val)
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); para._p.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None: sp = OxmlElement('w:spacing'); pPr.append(sp)
    for k, v in [('w:before', '0'), ('w:after', '0'), ('w:line', '168'), ('w:lineRule', 'auto')]:
        sp.set(qn(k), v)


def write_digit_cell(cell, digit):
    para = cell.paragraphs[0]
    for run in list(para.runs): run._element.getparent().remove(run._element)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(digit)
    run.font.size = Pt(8); run.font.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    # 나눔고딕 또는 Malgun Gothic (Docker 환경에서 나눔고딕 사용)
    font_name = 'NanumGothic'
    for attr in ['w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs']:
        rFonts.set(qn(attr), font_name)
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); para._p.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None: sp = OxmlElement('w:spacing'); pPr.append(sp)
    for k, v in [('w:before', '0'), ('w:after', '0'), ('w:line', '168'), ('w:lineRule', 'auto')]:
        sp.set(qn(k), v)


def get_unique_cells(row):
    seen, result = set(), []
    for i, cell in enumerate(row.cells):
        cid = id(cell._tc)
        if cid not in seen: seen.add(cid); result.append((i, cell))
    return result


def fill_form(data: dict) -> bytes:
    """
    data 딕셔너리 키:
        application_type: str  # registration/activity_permit/reissue/workplace_chg/
                               # extension/reentry/status_change/residence/
                               # status_grant/reg_change
        desired_status:   str  # 희망 체류자격 (예: E-7-4)
        surname:          str  # 성 (영문 대문자)
        given_names:      str  # 명 (영문 대문자)
        birth_year:       str  # 출생 연도 (4자리)
        birth_month:      str  # 출생 월 (2자리)
        birth_day:        str  # 출생 일 (2자리)
        sex:              str  # M 또는 F
        nationality:      str  # 국적 (영문 대문자)
        alien_reg_no:     str  # 외국인등록번호 (예: 920320-5123456)
        passport_no:      str  # 여권번호
        passport_issue:   str  # 여권 발급일 (YYYY-MM-DD)
        passport_expiry:  str  # 여권 유효기간 (YYYY-MM-DD)
        address_korea:    str  # 대한민국 내 주소
        phone:            str  # 전화번호
        cell_phone:       str  # 휴대전화
        address_home:     str  # 본국 주소
        phone_home:       str  # 본국 전화
        application_date: str  # 신청일 (예: 2026년 03월 15일)

    Returns:
        bytes: PDF 파일 내용 (1페이지)
    """
    doc = Document(TEMPLATE)
    table = doc.tables[0]

    # 행 높이 설정
    for row_idx, twips in ROW_HEIGHTS.items():
        if row_idx < len(table.rows):
            set_row_height(table.rows[row_idx], twips, exact=row_idx in EXACT_ROWS)

    # 행 17 (외국인등록번호) 레이블 폰트 축소
    seen = set()
    for cell in table.rows[17].cells:
        cid = id(cell._tc)
        if cid in seen: continue
        seen.add(cid)
        if cell.text.strip(): shrink_cell_english_font(cell, 14)

    # 행 24 (원근무처) 레이블 폰트 축소
    seen = set()
    for cell in table.rows[24].cells:
        cid = id(cell._tc)
        if cid in seen: continue
        seen.add(cid)
        if cell.text.strip(): shrink_cell_english_font(cell, 14)

    # 신청 유형 체크박스
    checkbox_map = {
        'registration': (4, 0), 'activity_permit': (4, 12),
        'reissue': (6, 0), 'workplace_chg': (6, 12),
        'extension': (7, 0), 'reentry': (7, 12),
        'status_change': (8, 0), 'residence': (8, 12),
        'status_grant': (10, 0), 'reg_change': (10, 12),
    }
    for code, (r, c) in checkbox_map.items():
        cell = table.rows[r].cells[c]
        for para in cell.paragraphs:
            for run in para.runs:
                if '[  ]' in run.text or '[ ]' in run.text:
                    if code == data.get('application_type', ''):
                        run.text = run.text.replace('[  ]', '[✓]').replace('[ ]', '[✓]')

    # 영문 성명 (행 14)
    seen = set()
    for c_idx, cell in enumerate(table.rows[14].cells):
        cid = id(cell._tc)
        if cid in seen: continue
        seen.add(cid)
        if c_idx == 2:   write_cell(cell, data.get('surname', ''))
        elif c_idx == 12: write_cell(cell, data.get('given_names', ''))

    # 생년월일 / 성별 / 국적 (행 15)
    seen = set()
    for c_idx, cell in enumerate(table.rows[15].cells):
        cid = id(cell._tc)
        if cid in seen: continue
        seen.add(cid)
        if c_idx == 3:    write_cell(cell, data.get('birth_year', ''))
        elif c_idx == 13: write_cell(cell, data.get('birth_month', ''))
        elif c_idx == 17: write_cell(cell, data.get('birth_day', ''))
        elif c_idx == 27:
            sex = data.get('sex', '')
            for para in cell.paragraphs:
                for run in para.runs:
                    if '남' in run.text or '여' in run.text:
                        if sex == 'M':
                            run.text = run.text.replace('[ ]남', '[✓]남').replace('[ ]여', '[ ]여')
                        elif sex == 'F':
                            run.text = run.text.replace('[ ]남', '[ ]남').replace('[ ]여', '[✓]여')
        elif c_idx == 38: write_cell(cell, data.get('nationality', ''))

    # 외국인등록번호 (행 17) - 한 자리씩 13칸
    unique17 = get_unique_cells(table.rows[17])
    label_kw = ['외국인등록번호', '국 적', '국적', 'Nationality', 'Foreign']
    reg_cells = [(c, cell) for c, cell in unique17
                 if not cell.text.strip()
                 and not any(k in cell.text for k in label_kw)
                 and c != 38]
    alien = data.get('alien_reg_no', '').replace('-', '').replace(' ', '')
    for i, (c, cell) in enumerate(reg_cells[:13]):
        if i < len(alien): write_digit_cell(cell, alien[i])

    # 여권번호 / 발급일 / 유효기간 (행 18)
    unique18 = get_unique_cells(table.rows[18])
    empty18 = [(c, cell) for c, cell in unique18 if not cell.text.strip() and c > 0]
    for i, val in enumerate([
        data.get('passport_no', ''),
        data.get('passport_issue', ''),
        data.get('passport_expiry', '')
    ]):
        if i < len(empty18): write_cell(empty18[i][1], val)

    # 대한민국 주소 (행 19)
    for c, cell in get_unique_cells(table.rows[19]):
        if not cell.text.strip() and c > 0:
            write_cell(cell, data.get('address_korea', ''))
            break

    # 전화번호 / 휴대전화 (행 20)
    phones20 = [(c, cell) for c, cell in get_unique_cells(table.rows[20])
                if not cell.text.strip() and c > 0]
    for i, val in enumerate([data.get('phone', ''), data.get('cell_phone', '')]):
        if i < len(phones20): write_cell(phones20[i][1], val)

    # 본국 주소 / 본국 전화 (행 21)
    empty21 = [(c, cell) for c, cell in get_unique_cells(table.rows[21])
               if not cell.text.strip() and c > 0]
    for i, val in enumerate([data.get('address_home', ''), data.get('phone_home', '')]):
        if i < len(empty21): write_cell(empty21[i][1], val)

    # 신청일 (행 29)
    for c, cell in get_unique_cells(table.rows[29]):
        if not cell.text.strip() and c > 0:
            write_cell(cell, data.get('application_date', ''))
            break

    # LibreOffice PDF 변환 → 1페이지만 추출
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'form.docx')
        doc.save(docx_path)

        # LibreOffice 동시 실행 충돌 방지: 독립 HOME 디렉토리
        env = os.environ.copy()
        env['HOME'] = tmpdir

        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmpdir, docx_path],
            capture_output=True, timeout=120, env=env
        )

        for f in os.listdir(tmpdir):
            if f.endswith('.pdf'):
                reader = PdfReader(os.path.join(tmpdir, f))
                writer = PdfWriter()
                writer.add_page(reader.pages[0])   # 1페이지만
                buf = io.BytesIO()
                writer.write(buf)
                return buf.getvalue()

    raise RuntimeError('LibreOffice PDF 변환 실패 - LibreOffice 설치 여부 확인')
